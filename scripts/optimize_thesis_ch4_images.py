from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.text.paragraph import Paragraph
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/thesis/毕业论文初稿.docx"
OUTPUT = ROOT / "docs/thesis/毕业论文初稿_第四章图片优化.docx"
PLAYWRIGHT = ROOT / "output/playwright"
LEGACY = ROOT / "docs/thesis/screenshots/ch4-pages-20260504"
OPT_DIR = ROOT / "docs/thesis/screenshots/ch4-optimized"


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def set_text(paragraph: Paragraph, text: str, style=None, align=None) -> None:
    clear_paragraph(paragraph)
    if style is not None:
        paragraph.style = style
    paragraph.add_run(text)
    if align is not None:
        paragraph.alignment = align


def set_picture(paragraph: Paragraph, image_path: Path, width_cm: float) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(optimized_image(image_path)), width=Cm(width_cm))


def insert_figure_block(after_para: Paragraph, intro: str, image_path: Path, width_cm: float, caption: str,
                        normal_style, image_style, caption_style) -> Paragraph:
    intro_para = insert_paragraph_after(after_para, intro, normal_style)
    image_para = insert_paragraph_after(intro_para, style=image_style)
    set_picture(image_para, image_path, width_cm)
    caption_para = insert_paragraph_after(image_para, caption, caption_style)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return caption_para


def optimized_image(source: Path, max_width: int = 1600, quality: int = 82) -> Path:
    OPT_DIR.mkdir(parents=True, exist_ok=True)
    target = OPT_DIR / f"{source.stem}.jpg"
    if target.exists() and target.stat().st_mtime >= source.stat().st_mtime:
        return target

    with Image.open(source) as img:
        img = img.convert("RGB")
        if img.width > max_width:
            ratio = max_width / img.width
            img = img.resize((max_width, int(img.height * ratio)), Image.LANCZOS)
        img.save(target, format="JPEG", quality=quality, optimize=True, progressive=True)
    return target


def main() -> None:
    doc = Document(SOURCE)
    paras = list(doc.paragraphs)

    normal_style = paras[166].style
    image_style = paras[167].style
    caption_style = paras[168].style
    alt_caption_style = paras[181].style
    table_caption_style = paras[274].style

    # 4.2 账号认证、角色入口与基础资料实现
    set_text(
        paras[166],
        "管理员后台在账号体系中承担治理角色。AdminServiceImpl 登录时会校验管理员状态、BCrypt 密码和角色权限，并为管理员生成包含权限信息的 JWT；RoleController、PermissionController 和 AdminOperationLogService 进一步支撑角色、权限和操作日志。这样做的意义在于，用户端和商家端负责业务动作，管理员后台负责确认谁可以审核、谁可以管理资金、谁可以查看系统日志，形成平台级访问边界。",
        style=normal_style,
    )
    intro_41 = insert_paragraph_after(paras[166], "用户端登录页面如图4.1所示。", normal_style)
    set_picture(paras[167], PLAYWRIGHT / "electron-user-login.png", 14.2)
    paras[168].alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro_42 = insert_paragraph_after(paras[168], "注册与找回密码入口如图4.2所示。", normal_style)
    set_text(intro_42, "注册与找回密码入口如图4.2所示。", style=normal_style)
    set_picture(paras[169], LEGACY / "fig4_9_login_actions.png", 11.5)
    paras[170].alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro_43 = insert_paragraph_after(paras[170], "管理员登录页面如图4.3所示。", normal_style)
    set_text(intro_43, "管理员登录页面如图4.3所示。", style=normal_style)
    set_picture(paras[171], PLAYWRIGHT / "electron-admin-login.png", 9.8)
    paras[172].alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro_44 = insert_paragraph_after(paras[172], "用户中心与基础资料页面如图4.4所示。", normal_style)
    set_text(intro_44, "用户中心与基础资料页面如图4.4所示。", style=normal_style)
    set_picture(paras[173], PLAYWRIGHT / "electron-user-profile.png", 14.6)
    paras[174].alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro_45a = insert_paragraph_after(paras[174], "地址管理页面如图4.5（a）所示。", normal_style)
    set_text(intro_45a, "地址管理页面如图4.5（a）所示。", style=normal_style)
    set_picture(paras[175], PLAYWRIGHT / "electron-user-address.png", 14.6)
    set_text(paras[176], "图4.5（a） 地址管理页面", style=caption_style, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap_45b = insert_figure_block(
        paras[176],
        "钱包管理页面如图4.5（b）所示。",
        PLAYWRIGHT / "electron-user-wallet-management.png",
        14.6,
        "图4.5（b） 钱包管理页面",
        normal_style,
        image_style,
        caption_style,
    )

    # 4.3 商家、菜单与菜品管理实现
    intro_46 = insert_paragraph_after(paras[179], "用户端商家查找页面如图4.6所示。", normal_style)
    set_picture(paras[180], PLAYWRIGHT / "electron-user-merchants.png", 14.8)
    paras[181].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paras[181].style = alt_caption_style

    intro_47 = insert_paragraph_after(paras[184], "商家端经营统计页面如图4.7所示。", normal_style)
    set_picture(paras[185], PLAYWRIGHT / "electron-merchant-statistics.png", 14.8)
    paras[186].alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_text(
        paras[187],
        "从图4.7可以看出，商家端已经从菜品与订单维护继续走向经营复盘。经营统计虽然不直接产生菜品数据，但它会反过来影响商家对菜单、菜品和活动的调整，因此也属于商家运营闭环的一部分。",
        style=normal_style,
    )
    intro_48 = insert_paragraph_after(paras[187], "商家端菜单管理页面如图4.8所示。", normal_style)
    set_picture(paras[188], PLAYWRIGHT / "electron-merchant-menu.png", 14.8)
    paras[189].alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro_49 = insert_paragraph_after(paras[189], "商家端菜品管理页面如图4.9所示。", normal_style)
    set_picture(paras[190], PLAYWRIGHT / "electron-merchant-dish-management.png", 14.8)
    paras[191].alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro_410a = insert_paragraph_after(paras[191], "用户端商家详情页面如图4.10（a）所示。", normal_style)
    set_picture(paras[192], PLAYWRIGHT / "electron-user-merchant-detail.png", 14.8)
    set_text(paras[193], "图4.10（a） 用户端商家详情页面", style=caption_style, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap_410b = insert_figure_block(
        paras[193],
        "用户端菜品详情页面如图4.10（b）所示。",
        PLAYWRIGHT / "electron-user-dish-detail.png",
        14.8,
        "图4.10（b） 用户端菜品详情页面",
        normal_style,
        image_style,
        caption_style,
    )

    # 4.4 个性化推荐、健康管理与食谱实现
    intro_412 = insert_paragraph_after(paras[201], "用户端卡路里统计页面如图4.12所示。", normal_style)
    set_picture(paras[202], ROOT / "论文插图/fig_4_4_user_calorie.png", 14.8)
    paras[203].alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_text(
        paras[205],
        "节日推荐和食谱能力进一步补充了推荐场景。FestivalController 提供当前节日、节日推荐和用户反馈入口，RecipeController 与 Tutorial 相关接口则承接食谱浏览、我的食谱和教程内容。管理员后台的数据统计、教程管理和内容审核为这些内容提供治理支撑，避免推荐和食谱只依赖临时页面配置。",
        style=normal_style,
    )
    intro_413 = insert_paragraph_after(paras[205], "用户端个性化推荐页面如图4.13所示。", normal_style)
    set_picture(paras[206], PLAYWRIGHT / "electron-user-recommend.png", 14.8)
    paras[207].alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro_414a = insert_paragraph_after(paras[207], "用户端我的食谱页面如图4.14（a）所示。", normal_style)
    set_picture(paras[208], PLAYWRIGHT / "electron-user-my-recipe.png", 14.8)
    set_text(paras[209], "图4.14（a） 用户端我的食谱页面", style=caption_style, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap_414b = insert_figure_block(
        paras[209],
        "用户端饮食记录页面如图4.14（b）所示。",
        PLAYWRIGHT / "electron-user-diet-record.png",
        14.8,
        "图4.14（b） 用户端饮食记录页面",
        normal_style,
        image_style,
        caption_style,
    )

    # 4.5 订单、支付、钱包与拼单实现
    set_text(
        paras[219],
        "桌面端用户侧覆盖订单确认、我的订单、订单详情、评价订单和钱包管理，商家侧覆盖订单管理、今日订单和订单详情；小程序端覆盖购物车、确认订单、订单详情、订单进度、我的订单、拼单创建、加入、选菜和结算。管理员后台的订单管理、退款管理、充值记录和提现审核则承担交易风险收口。",
        style=normal_style,
    )
    intro_416 = insert_paragraph_after(paras[219], "商家端订单管理页面如图4.16所示。", normal_style)
    set_picture(paras[220], PLAYWRIGHT / "electron-merchant-orders.png", 14.8)
    paras[221].alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_text(
        paras[222],
        "从图4.16可以看出，商家端订单管理页面支持商家在同一工作台中查看订单列表、处理订单状态和进入订单详情，后端订单状态变化又会通过消息和统计链路影响用户端订单进度、商家今日订单和后台订单管理。",
        style=normal_style,
    )
    intro_417a = insert_paragraph_after(paras[222], "用户端我的订单页面如图4.17（a）所示。", normal_style)
    set_picture(paras[223], PLAYWRIGHT / "electron-user-orders.png", 14.8)
    set_text(paras[224], "图4.17（a） 用户端我的订单页面", style=caption_style, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap_417b = insert_figure_block(
        paras[224],
        "用户端订单详情页面如图4.17（b）所示。",
        PLAYWRIGHT / "electron-user-order-detail.png",
        14.8,
        "图4.17（b） 用户端订单详情页面",
        normal_style,
        image_style,
        caption_style,
    )

    intro_418a = insert_paragraph_after(cap_417b, "用户端确认订单页面如图4.18（a）所示。", normal_style)
    set_picture(paras[225], PLAYWRIGHT / "electron-before-pay.png", 14.8)
    set_text(paras[226], "图4.18（a） 用户端确认订单页面", style=caption_style, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap_418b = insert_figure_block(
        paras[226],
        "支付确认弹窗如图4.18（b）所示。",
        PLAYWRIGHT / "electron-user-payment-confirm-dialog.png",
        14.8,
        "图4.18（b） 支付确认弹窗",
        normal_style,
        image_style,
        caption_style,
    )

    intro_419a = insert_paragraph_after(cap_418b, "拼单创建页面如图4.19（a）所示。", normal_style)
    set_picture(paras[227], PLAYWRIGHT / "electron-grouporder-drawer.png", 14.8)
    set_text(paras[228], "图4.19（a） 拼单创建页面", style=caption_style, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap_419b = insert_figure_block(
        paras[228],
        "拼单结算页面如图4.19（b）所示。",
        PLAYWRIGHT / "electron-group-order-drawer.png",
        14.8,
        "图4.19（b） 拼单结算页面",
        normal_style,
        image_style,
        caption_style,
    )

    # 4.6 消息、社交协同与通知实现
    set_text(
        paras[235],
        "通知和公告承担平台支撑作用。NotificationController 负责用户通知列表、未读数、已读和删除；AdminAnnouncementController 负责后台公告管理；MessageRecordController 提供消息记录、未读数、发送和批量删除。管理员后台还可以通过系统日志观察异常行为，形成消息链路的治理闭环。",
        style=normal_style,
    )
    intro_420a = insert_paragraph_after(paras[235], "用户端消息中心页面如图4.20（a）所示。", normal_style)
    set_picture(paras[236], PLAYWRIGHT / "electron-user-message-center.png", 14.8)
    set_text(paras[237], "图4.20（a） 用户端消息中心页面", style=caption_style, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap_420b = insert_figure_block(
        paras[237],
        "用户端聊天页面如图4.20（b）所示。",
        PLAYWRIGHT / "electron-user-chat.png",
        14.8,
        "图4.20（b） 用户端聊天页面",
        normal_style,
        image_style,
        caption_style,
    )

    # 4.8 教程、心愿单与运营内容实现
    set_text(
        paras[258],
        "热点话题、公告、优惠券、节日推荐和帮助反馈属于运营支撑内容。HotTopicAdminController、AdminAnnouncementController、CouponController、FestivalController 和用户反馈接口共同支撑首页内容、节日活动、优惠券校验、公告发布和意见反馈。小程序端配置了食谱、心愿单、帮助、反馈和优惠券页面，桌面端也配置了热点、收藏、教程和心愿单审核入口。这些模块不直接产生订单，但会影响用户发现内容、商家响应需求和平台维护内容质量。",
        style=normal_style,
    )
    intro_424a = insert_paragraph_after(paras[258], "用户端教程广场页面如图4.24（a）所示。", normal_style)
    set_picture(paras[259], PLAYWRIGHT / "electron-user-tutorials.png", 14.8)
    set_text(paras[260], "图4.24（a） 用户端教程广场页面", style=caption_style, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap_424b = insert_figure_block(
        paras[260],
        "用户端我的教程页面如图4.24（b）所示。",
        PLAYWRIGHT / "electron-user-my-tutorials.png",
        14.8,
        "图4.24（b） 用户端我的教程页面",
        normal_style,
        image_style,
        caption_style,
    )

    intro_425a = insert_paragraph_after(cap_424b, "商家端心愿单审核页面如图4.25（a）所示。", normal_style)
    set_picture(paras[261], PLAYWRIGHT / "electron-merchant-wish-list-audit.png", 14.8)
    set_text(paras[262], "图4.25（a） 商家端心愿单审核页面", style=caption_style, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap_425b = insert_figure_block(
        paras[262],
        "商家端教程管理页面如图4.25（b）所示。",
        PLAYWRIGHT / "electron-merchant-tutorials.png",
        14.8,
        "图4.25（b） 商家端教程管理页面",
        normal_style,
        image_style,
        caption_style,
    )

    # 4.9 管理员后台实现
    set_text(
        paras[268],
        "运维统计包括数据统计、系统日志、缓存监控和 Agent 监控。AdminStatisticsController 提供后台统计视图，SystemLogController 和 AdminOperationLogService 记录系统与管理员操作，CacheMonitorController 提供缓存状态、详情、报告和健康检查，AgentMonitoringController 提供 Agent 性能、调用链和概览。管理员后台因此承担了安全、内容质量、交易秩序和运行状态四类支撑职责。",
        style=normal_style,
    )
    intro_426a = insert_paragraph_after(paras[268], "管理员登录页面如图4.26（a）所示。", normal_style)
    set_picture(paras[269], PLAYWRIGHT / "electron-admin-login.png", 9.8)
    set_text(paras[270], "图4.26（a） 管理员登录页面", style=caption_style, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap_426b = insert_figure_block(
        paras[270],
        "管理员控制台页面如图4.26（b）所示。",
        PLAYWRIGHT / "electron-admin-dashboard.png",
        14.8,
        "图4.26（b） 管理员控制台页面",
        normal_style,
        image_style,
        caption_style,
    )

    intro_427a = insert_paragraph_after(cap_426b, "管理员商家审核页面如图4.27（a）所示。", normal_style)
    set_picture(paras[271], PLAYWRIGHT / "electron-admin-merchants-audit.png", 14.8)
    set_text(paras[272], "图4.27（a） 管理员商家审核页面", style=caption_style, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap_427b = insert_figure_block(
        paras[272],
        "管理员菜品审核页面如图4.27（b）所示。",
        PLAYWRIGHT / "electron-admin-dishes-audit.png",
        14.8,
        "图4.27（b） 管理员菜品审核页面",
        normal_style,
        image_style,
        caption_style,
    )

    intro_428a = insert_paragraph_after(cap_427b, "管理员订单管理页面如图4.28（a）所示。", normal_style)
    set_picture(paras[273], PLAYWRIGHT / "electron-admin-orders.png", 14.8)
    set_text(paras[274], "图4.28（a） 管理员订单管理页面", style=table_caption_style, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap_428b = insert_figure_block(
        paras[274],
        "管理员退款管理页面如图4.28（b）所示。",
        PLAYWRIGHT / "electron-admin-refunds.png",
        14.8,
        "图4.28（b） 管理员退款管理页面",
        normal_style,
        image_style,
        caption_style,
    )
    cap_428c = insert_figure_block(
        cap_428b,
        "管理员充值记录页面如图4.28（c）所示。",
        PLAYWRIGHT / "electron-admin-recharges.png",
        14.8,
        "图4.28（c） 管理员充值记录页面",
        normal_style,
        image_style,
        caption_style,
    )

    intro_429a = insert_paragraph_after(cap_428c, "管理员权限管理页面如图4.29（a）所示。", normal_style)
    set_picture(paras[275], PLAYWRIGHT / "electron-admin-permissions.png", 14.8)
    set_text(paras[276], "图4.29（a） 管理员权限管理页面", style=caption_style, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap_429b = insert_figure_block(
        paras[276],
        "管理员系统日志页面如图4.29（b）所示。",
        PLAYWRIGHT / "electron-admin-logs.png",
        14.8,
        "图4.29（b） 管理员系统日志页面",
        normal_style,
        image_style,
        caption_style,
    )
    insert_figure_block(
        cap_429b,
        "管理员数据统计页面如图4.29（c）所示。",
        PLAYWRIGHT / "electron-admin-statistics.png",
        14.8,
        "图4.29（c） 管理员数据统计页面",
        normal_style,
        image_style,
        caption_style,
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
